"""
WC KB — Document Chunker
Extracts text from PDF/DOCX/TXT and splits into overlapping chunks.
"""
from __future__ import annotations
import re, logging
from pathlib import Path
from typing import List

log = logging.getLogger("wc.kb.chunker")

CHUNK_SIZE    = int(__import__("os").getenv("CHUNK_SIZE",    "512"))   # words
CHUNK_OVERLAP = int(__import__("os").getenv("CHUNK_OVERLAP", "50"))    # words


def extract_text(file_path: str | Path) -> str:
    """Extract plain text from PDF, DOCX, or TXT file."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path)
    elif suffix == ".docx":
        return _extract_docx(path)
    elif suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _extract_pdf(path: Path) -> str:
    """Extract text from PDF. Falls back to OCR if no text layer."""
    text = ""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages  = []
        for page in reader.pages:
            t = page.extract_text() or ""
            pages.append(t)
        text = "\n".join(pages).strip()
    except Exception as e:
        log.warning(f"[chunker] pypdf failed for {path.name}: {e}")

    # If no text extracted, try pdfplumber
    if not text:
        try:
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                text = "\n".join(
                    (p.extract_text() or "") for p in pdf.pages
                ).strip()
        except Exception as e:
            log.warning(f"[chunker] pdfplumber failed for {path.name}: {e}")

    if not text:
        raise ValueError(f"Could not extract text from {path.name} — no text layer found.")

    return text


def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def chunk_text(
    text: str,
    doc_name: str,
    chunk_size: int = CHUNK_SIZE,
    overlap: int = CHUNK_OVERLAP,
) -> List[dict]:
    """
    Split text into overlapping word-based chunks.
    Returns list of dicts with: text, doc_name, chunk_index, page_num (estimated).
    """
    # Clean up whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)

    # Split into sentences first for cleaner boundaries
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    sentences = [s.strip() for s in sentences if s.strip()]

    chunks   = []
    current  = []
    cur_words = 0
    idx      = 0

    for sent in sentences:
        words = sent.split()
        w_count = len(words)

        if cur_words + w_count > chunk_size and current:
            # Save current chunk
            chunk_text = " ".join(current)
            if len(chunk_text.strip()) > 30:  # skip tiny chunks
                chunks.append({
                    "text":        chunk_text.strip(),
                    "doc_name":    doc_name,
                    "chunk_index": idx,
                    "page_num":    _estimate_page(idx, len(chunks)),
                })
                idx += 1

            # Overlap: keep last N words
            overlap_words = " ".join(current).split()[-overlap:]
            current   = overlap_words + words
            cur_words = len(current)
        else:
            current.extend(words)
            cur_words += w_count

    # Final chunk
    if current:
        chunk_text = " ".join(current)
        if len(chunk_text.strip()) > 30:
            chunks.append({
                "text":        chunk_text.strip(),
                "doc_name":    doc_name,
                "chunk_index": idx,
                "page_num":    _estimate_page(idx, len(chunks)),
            })

    log.info(f"[chunker] {doc_name}: {len(chunks)} chunks from {len(text.split())} words")
    return chunks


def _estimate_page(chunk_idx: int, total_chunks: int) -> int:
    """Rough page number estimate (assumes ~5 chunks per page)."""
    return (chunk_idx // 5) + 1
